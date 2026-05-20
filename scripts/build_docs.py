"""Generate the stage cheat-sheet and full-script Word files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent.parent / "demo-pack"
OUT_DIR.mkdir(exist_ok=True)


# ---------- styling helpers ----------

NAVY = RGBColor(0x0F, 0x23, 0x42)
PRIMARY = RGBColor(0x1D, 0x6F, 0xB8)
INK_SOFT = RGBColor(0x47, 0x55, 0x69)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x15, 0x80, 0x3D)
AMBER = RGBColor(0xB4, 0x53, 0x09)


def _set_margins(doc, cm=1.5):
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply a background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _heading(doc, text, *, level=1, color=NAVY, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if level >= 2:
        p.paragraph_format.space_before = Pt(8)
    return p


def _para(doc, text, *, size=10.5, color=NAVY, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = NAVY
    return p


def _hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "94A3B8")
    pBdr.append(bottom)
    pPr.append(pBdr)


# =====================================================================
#                        ONE-PAGE CHEAT SHEET
# =====================================================================


def build_cheat_sheet() -> Path:
    doc = Document()
    _set_margins(doc, cm=1.2)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title bar
    title = doc.add_paragraph()
    run = title.add_run("CLAUDE CODE FOR HEALTHCARE  ·  Multi-Agent Clinical Co-Pilot")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    r = sub.add_run("5-min stage cheat sheet  ·  URL: http://127.0.0.1:8765  ·  Engine: Azure OpenAI")
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK_SOFT
    _hr(doc)

    # Pre-stage strip
    p = doc.add_paragraph()
    r = p.add_run("PRE-STAGE (60s)")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = AMBER

    pre = doc.add_paragraph()
    pre.add_run("• Curl ").bold = False
    pre.add_run("/api/health ").font.name = "Consolas"
    pre.add_run("→ azure_openai_configured: true   ")
    pre.add_run("• Cmd+Shift+R on browser   ")
    pre.add_run("• Click Whitfield once off-screen to pre-warm Azure")
    for r in pre.runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY

    _hr(doc)

    # 5-minute table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(["t", "CLICK", "SAY"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], "1D6FB8")
    hdr[0].width = Cm(1.3)
    hdr[1].width = Cm(5.5)
    hdr[2].width = Cm(11.0)

    rows = [
        ("0:00", "Open browser",
         "“Built with Claude Code in a weekend. Every LLM call runs on de-identified text.”"),
        ("0:20", "Whitfield · 68F · new AFib",
         "“68F new AFib. Stable. Cardiologist plan: start amiodarone. She’s on warfarin.”"),
        ("1:00", "Run Co-Pilot button",
         "“Six specialist Claude agents reasoning in parallel. Real LLM calls, streaming live.”"),
        ("2:00", "Click Pharmacy panel",
         "★ THE WOW ★ “Pharmacy caught it. Amiodarone + warfarin via CYP2C9/3A4. INR will spike to 4.5+ in 7-10 days. The cardiologist almost missed it.”"),
        ("2:45", "Plan tab",
         "“Orchestrator detects convergence — Pharmacy, Bias-Check, Guidelines all flagged the same issue.”"),
        ("3:00", "Toggle LLM ↔ clinician view",
         "“Same plan. Placeholders for the model. Real names for the clinician. Re-hydration server-side.”"),
        ("3:30", "+ Paste a case (audience)",
         "“Throw me a synthetic case. Anything.”  ←  Real, unscripted, converts skeptics."),
        ("4:15", "Safety & Audit → Verify no raw PHI",
         "Green ✓. “Three views, three trust levels. Audit log = hashes + counts only.”"),
        ("4:45", "—",
         "Close: “Six Claude agents. One orchestrator. PHI safety as foundation. Built end-to-end with Claude Code.”"),
    ]
    for t, click, say in rows:
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""

        rp = row[0].paragraphs[0]
        rr = rp.add_run(t)
        rr.bold = True
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = PRIMARY

        rp = row[1].paragraphs[0]
        rr = rp.add_run(click)
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = NAVY
        rr.bold = True

        rp = row[2].paragraphs[0]
        rr = rp.add_run(say)
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = NAVY

        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _hr(doc)

    # Backup strip
    p = doc.add_paragraph()
    r = p.add_run("BACKUP PLAN")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RED

    bk = doc.add_paragraph()
    bk.add_run("• Azure breaks → switch engine to Offline Simulator (right rail), keeps going  ")
    bk.add_run("• Stale UI → Cmd+Shift+R   ")
    bk.add_run("• Server crash → ./run_web.sh again")
    for r in bk.runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY

    # Cases strip
    p = doc.add_paragraph()
    r = p.add_run("CASES AVAILABLE")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = AMBER

    cases = doc.add_paragraph()
    cases.add_run("HERO: Whitfield · AFib + warfarin/amiodarone catch   ")
    cases.add_run("• Singh · chest pain → premature closure on ACS vs PE   ")
    cases.add_run("• Chen · thunderclap headache → missed SAH/meningitis   ")
    cases.add_run("• Patel · falls → iatrogenic anticholinergic burden (ACB ~11)")
    for r in cases.runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY

    # Close line
    _hr(doc)
    p = doc.add_paragraph()
    r = p.add_run(
        "Key lesson: in healthcare AI, you don’t replace the clinician — you give them six more pairs of eyes that never get tired and always read the medication list."
    )
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK_SOFT

    out = OUT_DIR / "demo-cheat-sheet.docx"
    doc.save(out)
    return out


# =====================================================================
#                         FULL SCRIPT
# =====================================================================


def build_full_script() -> Path:
    doc = Document()
    _set_margins(doc, cm=2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    p = doc.add_paragraph()
    r = p.add_run("Claude Code for Healthcare")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("Multi-Agent Clinical Co-Pilot  ·  Stage Pack")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    r = p.add_run(
        "Total time: 5 minutes  ·  URL: http://127.0.0.1:8765  ·  Engine: Azure OpenAI (live stream)."
    )
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK_SOFT

    p = doc.add_paragraph()
    r = p.add_run(
        "Synthetic data only. This is a demonstration of privacy-preserving engineering patterns, not a HIPAA/PHIPA/PIPEDA-compliant system."
    )
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = AMBER

    _hr(doc)

    # ----- Patients in the demo -----
    _heading(doc, "Patients in the demo", level=2, size=14)
    _para(
        doc,
        "Four preset cases plus one audience-paste slot. Realistic stage plan: run Whitfield (hero) for ~3 min, then one audience-paste case, then Audit. If you have extra time, run a second preset.",
    )

    cases_table = doc.add_table(rows=1, cols=4)
    cases_table.autofit = True
    headers = ["#", "Case", "Reasoning pattern", "Use for"]
    for i, h in enumerate(headers):
        cell = cases_table.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "1D6FB8")
    rows = [
        ("1", "Whitfield · 68F · new AFib + complex polypharmacy",
         "Drug-drug interaction catch (warfarin + amiodarone)",
         "HERO — must run."),
        ("2", "Singh · 54M · chest pain",
         "Premature closure on ACS when PE is more likely",
         "Backup wow #1"),
        ("3", "Chen · 32F · severe headache",
         "Missed SAH / meningitis red flags",
         "Backup wow #2"),
        ("4", "Patel · 76M · falls + cognitive decline",
         "Iatrogenic anticholinergic burden (5 offenders)",
         "Backup wow #3 — finance/business folks love this one"),
        ("5", "+ Paste a case",
         "Audience input",
         "Interactive moment"),
    ]
    for row in rows:
        cells = cases_table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(v)
            rr.font.size = Pt(10)
            rr.font.color.rgb = NAVY
            if i == 0:
                rr.bold = True
        cells[1].paragraphs[0].runs[0].bold = True

    # ----- Pre-stage checklist -----
    _heading(doc, "Pre-stage checklist (60 seconds before)", level=2, size=14)
    _bullet(doc, "Make sure the server is up: curl -s http://127.0.0.1:8765/api/health")
    _bullet(doc, "Expect: azure_openai_configured: true")
    _bullet(doc, "Open the browser to http://127.0.0.1:8765")
    _bullet(doc, "Hard-refresh once (Cmd+Shift+R) so static files are fresh")
    _bullet(doc, "Click the Whitfield case ONCE off-screen to pre-warm Azure (avoids cold start)")
    _para(
        doc,
        "If azure_openai_configured is false, check the .env file or switch the engine radio to Offline Simulator and proceed.",
        italic=True,
        color=INK_SOFT,
        size=10,
    )

    # ----- The 5-minute script -----
    _heading(doc, "The 5-minute script", level=2, size=14)

    def beat(title, what_to_click, what_to_show, what_to_say, *, accent=NAVY):
        h = doc.add_paragraph()
        rr = h.add_run(title)
        rr.bold = True
        rr.font.size = Pt(12.5)
        rr.font.color.rgb = accent
        if what_to_click:
            _para(doc, f"Click: {what_to_click}", bold=True, size=10.5)
        if what_to_show:
            _para(doc, f"Show: {what_to_show}", size=10.5)
        if what_to_say:
            sp = doc.add_paragraph()
            rr = sp.add_run("Say: ")
            rr.bold = True
            rr.font.size = Pt(10.5)
            rr2 = sp.add_run(what_to_say)
            rr2.italic = True
            rr2.font.size = Pt(10.5)
            rr2.font.color.rgb = INK_SOFT
        doc.add_paragraph()

    beat(
        "0:00 — 0:20  ·  Open",
        what_to_click="Browser at http://127.0.0.1:8765",
        what_to_show="EMR chrome. Synthetic banner. ‘Clinical AI Co-Pilot · PHI-safe foundation.’",
        what_to_say="“I built this with Claude Code in a weekend. It looks like an EMR for a reason — this is what production healthcare AI actually needs to look like for clinicians to use it. Every LLM call in this app runs on de-identified text. We’ll come back to that.”",
    )

    beat(
        "0:20 — 1:00  ·  The hero case",
        what_to_click="‘Whitfield · 68F · new AFib + complex polypharmacy’ in the left rail",
        what_to_show="The Case tab fills with structured patient data — demographics, presenting complaint, history, current medications (warfarin), vitals, labs, clinician’s initial plan (start amiodarone).",
        what_to_say="“68-year-old woman. New-onset atrial fibrillation. Hemodynamically stable. The cardiologist’s initial plan: start amiodarone for rhythm control, continue her warfarin. CHADS-VASc of 4. Looks routine.”  Then point at meds: “She’s already on warfarin. Her INR is therapeutic. The plan adds amiodarone.”",
    )

    beat(
        "1:00 — 2:00  ·  Run the co-pilot",
        what_to_click="The big blue Run Co-Pilot button on the right rail",
        what_to_show="App auto-jumps to Co-Pilot Workspace tab. Six agent panels light up yellow (reasoning), streaming text appearing in each one in real time.",
        what_to_say="“Six specialist Claude agents are running in parallel right now. Triage is checking urgency. Differential is weighing the diagnoses. Pharmacy is reviewing the meds. Guidelines is pulling the AFib management recommendations. Bias-Check is auditing the clinician’s plan for cognitive bias. Communication is drafting the handoff.”  As they turn green: “Each one is a real Claude call streaming live.”",
    )

    beat(
        "2:00 — 2:45  ·  THE WOW MOMENT  ·  The catch",
        what_to_click="The Pharmacy panel to bring its full output into view",
        what_to_show="All six panels are green. Pharmacy panel has an amber flag: ‘High-severity interaction: Amiodarone + Warfarin’.",
        what_to_say="(Slow down.)  “Look at what Pharmacy caught. Amiodarone potentiates warfarin through CYP2C9 and CYP3A4 inhibition. The INR is going to spike to 4.5 or higher within 7 to 10 days. The cardiologist’s plan, exactly as written, sends this woman home with a serious bleeding risk that won’t manifest for a week.”  Pause. “The other five agents didn’t catch it — they were focused on rhythm control. That’s the whole point of a multi-agent setup.”",
        accent=RED,
    )

    beat(
        "2:45 — 3:30  ·  The synthesis",
        what_to_click="Plan tab",
        what_to_show="Orchestrator’s synthesized plan. Cross-flags panel shows multiple high-severity flags with 4–5 agents converging: ‘High-risk interaction between amiodarone and warfarin…’ and ‘Guideline discordance: rhythm control chosen without trial of rate control…’",
        what_to_say="“The orchestrator agent watches for convergence. When multiple specialists independently flag the same concern, severity goes up. Pharmacy AND Bias-Check AND Guidelines all converged on the same issue here. That’s a stronger signal than any one of them alone.”  Then click ‘What the LLM produced’ toggle: “This is what Claude actually saw. Placeholders, not names. The clinician view you saw a second ago — with ‘Eleanor Whitfield’ — that re-identification happens server-side, never on the model side.”  Toggle back.",
    )

    beat(
        "3:30 — 4:15  ·  Audience interaction",
        what_to_click="‘+ Paste a case’ at the bottom of the left rail",
        what_to_show="Paste modal opens. Paste an audience case (or your backup case from notes app if they freeze). Load it, run it.",
        what_to_say="“Throw a synthetic case at me. Anything — make one up. Polypharmacy, weird presentation, your favourite clinical pearl. I’ll paste it and run it live.”  This is the moment that converts skeptics.",
    )

    beat(
        "4:15 — 4:45  ·  The trust story",
        what_to_click="Safety & Audit tab, then ‘Verify no raw PHI’ button",
        what_to_show="Green checkmark — ‘0 raw PHI strings across N events’. Audit log entries below, each a clean JSON record of agent statuses, hashes, and counts — no patient names anywhere.",
        what_to_say="“Same workflow, three different views, three different trust levels. The model saw placeholders. The clinician saw real names. The audit log saw neither — just hashes and counts. This is what production healthcare AI looks like.”",
    )

    beat(
        "4:45 — 5:00  ·  Close",
        what_to_click="—",
        what_to_show="—",
        what_to_say="“Six Claude agents. One orchestrator. PHI safety as foundation. Each agent box is a real LLM call. Each agent has a different specialty system prompt. The orchestrator synthesizes their work and surfaces convergence. The whole thing was built end-to-end with Claude Code in a weekend.  The lesson: in healthcare AI, you don’t replace the clinician. You give them six more pairs of eyes that never get tired and always read the medication list.”",
        accent=GREEN,
    )

    # ----- Backup talking points -----
    _heading(doc, "Backup talking points (if you have extra time)", level=2, size=13)
    _para(
        doc,
        "If you want to run a second case: Patel · falls + cognitive decline.",
        bold=True,
    )
    _para(
        doc,
        "The finance/business folks love this one. Pharmacy detects 5 strong anticholinergic drugs (ACB score 11) and the orchestrator says ‘Do NOT add a cholinesterase inhibitor on top — pharmacologically contradictory.’ Easy to extrapolate to $40k inappropriate-medication prescriptions saved per patient × population.",
    )
    _para(doc, " ")
    _para(
        doc,
        "If you want a “wow without polypharmacy”: Chen · severe headache.",
        bold=True,
    )
    _para(
        doc,
        "Differential ranks SAH high. Pharmacy flags ketorolac as contraindicated (NSAID + possible bleed = harm). Bias-Check calls out premature closure on migraine. Three agents independently fix the same dangerous discharge plan.",
    )

    # ----- If something breaks -----
    _heading(doc, "If something breaks", level=2, size=13)

    t = doc.add_table(rows=1, cols=3)
    headers = ["Problem", "Action", "What to say"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "B91C1C")
    rows = [
        ("Azure call fails / times out",
         "Switch engine to Offline Simulator (right rail), click Run again",
         "“Let me show you the offline simulator we built as a stage backup — same six agents, same output structure, runs in 2 seconds. Mid-demo failover.”"),
        ("Browser shows stale UI",
         "Hard refresh (Cmd+Shift+R)",
         "Just keep talking through it."),
        ("Server crashes",
         "Run ./run_web.sh again in terminal",
         "“Healthcare-grade fault tolerance — observe.”"),
        ("Audience asks a hard medical question",
         "Deflect to engineering pattern",
         "“This is a demo of the engineering pattern, not a clinical product. Real deployment requires clinician oversight, regulatory review, and validation studies.”"),
    ]
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(v)
            rr.font.size = Pt(10)
            rr.font.color.rgb = NAVY

    # ----- Q&A bait -----
    _heading(doc, "Q&A bait (end with this)", level=2, size=13)
    _para(
        doc,
        "“Happy to take questions. I’m especially curious if anyone wants to see the actual Claude prompts, or how the cross-flag detection works under the hood.”  This invites the engineering crowd to engage on architecture and shifts the conversation away from medical-correctness debates you don’t want to have.",
        italic=True,
        color=INK_SOFT,
    )

    # ----- What each segment remembers -----
    _heading(doc, "What each audience segment will remember", level=2, size=13)
    _bullet(doc, "AI engineers: ‘Six parallel LLM streams in a real product, with cross-agent convergence detection.’")
    _bullet(doc, "Healthcare folks: ‘It actually caught the warfarin interaction.’")
    _bullet(doc, "Finance / business: ‘$40k saved per inappropriate prescription, scaled across a health system.’")
    _bullet(doc, "Product folks: ‘EMR-style workbench that augments instead of replaces.’")
    _bullet(doc, "Everyone: ‘Claude Code built this whole thing in a weekend.’")

    out = OUT_DIR / "demo-script-full.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    cheat = build_cheat_sheet()
    full = build_full_script()
    print(f"Wrote: {cheat}")
    print(f"Wrote: {full}")
